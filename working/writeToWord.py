import docx

def writeToDoc(test):
    doc=docx.Document()
    doc.add_paragraph(test)
    doc.add_picture('testPhoto.jpg',width=docx.shared.Inches(4),height=docx.shared.Inches(3))
    doc.save('new.docx')